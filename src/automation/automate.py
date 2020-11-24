"""
Main worker class for running the automation process.
"""
import logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger('email-automation')

import os
import traceback
import tempfile
import re
from docx import Document

from .parse import read_tables, order_vars, parse_if_expression, parse_value_expression, parse_format_expression
from .typography import unite_similar_runs, insert_runs
from .data import read_data

from PyQt5.QtCore import QObject, pyqtSignal


class StopException(Exception):
    """
    This exception is raised when the automation process needs to stop running or is finished.
    It doesn't actually do anything, instead we catch it in Automation.run() in order to close the automation thread.
    """
    pass


class GuiUpdater(logging.Handler, QObject):
    """
    Logging handler class that will update the GUI messages in another thread
    """
    logging_signal = pyqtSignal(int, str, int)  # Level, message, progress

    def __init__(self, has_gui=True):
        super().__init__()
        QObject.__init__(self)
        self.has_gui = has_gui

    def emit(self, record):
        """
        Emits a record via PyQt signal, if GUI is enabled, otherwise just print to console.

        Parameters
        ----------
        record : logging.LogRecord
            logging object that contains the logging message and the message's level and any args that were passed in.
        """

        if not self.has_gui:
            return

        if record.args:
            progress = record.args[0]
            record.args = None  # Probably not the best way to do this, but it'll work
        else:
            progress = None

        msg = self.format(record)
        self.logging_signal.emit(record.levelno, msg, progress)
        if record.levelno == logging.CRITICAL:
            raise StopException


class Automate(QObject):
    """
    Main worker object class intended to run in a QThread.
    It will run the automation process and report back to the GUI with status updates, warnings, and progress.
    """
    finished_signal = pyqtSignal()

    def __init__(self, in_gui=True):
        super().__init__()
        self.in_gui = in_gui
        self.guiUpdater = GuiUpdater(in_gui)
        self.filled_template = None

    def run(self, template_file, vars_file):
        """
        This function is called from the GUI thread. It starts the logging process and will catch any exceptions
        that are not caught during the automation process.

        Parameters
        ----------
        template_file : str
            The file path for the template document.
        vars_file : str
            The file path for the vars spreadsheet.
        """
        if logger.handlers:
            logger.removeHandler(logger.handlers[0])
        logger.addHandler(self.guiUpdater)

        try:
            self.automate(template_file, vars_file)

        except StopException:
            self.thread().quit()
            return
        except Exception as e:
            tb = traceback.format_exc()
            logger.error(f'An unexpected error has occurred: \n\n{tb}')
            self.thread().quit()
            return


    def automate(self, template_file, vars_file):
        """
        Orchestrates the automation process.

        Parameters
        ----------
        template_file : str
            The file path for the template document.
        vars_file : str
            The file path for the vars spreadsheet.
        """

        logger.info('Reading variables workbook and validating tables', 10)
        files, queries, sources, values, formats = read_tables(vars_file)


        logger.info('Getting initial values for variables (this could take a while...)')
        variables = read_data(files, queries, sources)


        logger.info('Evaluating variables', 50)
        for var in reversed(order_vars(values)):
            if var not in variables:  # Catches possible duplicates in the stack (can this even happen anymore?)
                for i in range(1, (len(values.columns) - 2) // 2 + 1):
                    if parse_if_expression(var, f'If{i}', 'Values', values.loc[var][f'If{i}'], variables):
                        expression = values.loc[var][f'Value{i}']
                        variables[var] = parse_value_expression(var, f'Value{i}', expression, variables)
                        break
                else:
                    expression = values.loc[var]['ValueElse']
                    variables[var] = parse_value_expression(var, 'ValueElse', expression, variables)

        logger.info('Evaluating variable formatting', 70)
        formatted_variables = {}
        for var in formats.index:
            for i in range(1, (len(formats.columns) - 1) // 2 + 1):
                if parse_if_expression(var, f'If{i}', 'Formats', formats.loc[var][f'If{i}'], variables):
                    formatted_variables[var] = parse_format_expression(var, f'Format{i}', formats.loc[var][f'Format{i}'], variables)
                    break
            else:
                formatted_variables[var] = parse_format_expression(var, 'FormatElse', formats.loc[var]['FormatElse'], variables)



        logger.info('Inserting variables into template', 90)
        template = Document(template_file)
        table_paragraphs = [paragraph for table in template.tables for row in table.rows
                            for cell in row.cells for paragraph in cell.paragraphs]
        paragraphs = template.paragraphs + table_paragraphs

        for paragraph in paragraphs:
            unite_similar_runs(paragraph.runs)
            i = 0
            while i < len(paragraph.runs):
                run = paragraph.runs[i]
                match = re.search(r"`(.+?)`", run.text)

                if match:
                    var = match.group()[1:-1]
                    if var in formatted_variables:
                        insert_runs(paragraph, formatted_variables[var], i, match.span())
                    elif var in variables:
                        logger.warning(f'No formatting specified for variable "{var}", it will have default font')
                        temporary_formatted_runs = [Document().add_paragraph().add_run(str(variables[var]))]
                        insert_runs(paragraph, temporary_formatted_runs, i, match.span())
                    else:
                        logger.warning(f'Template contains variable "{var}" that is not defined')
                i += 1

        with tempfile.TemporaryFile(suffix='.docx', delete=False) as f:
            template.save(f)

        os.startfile(f.name)  # Opens the filled report in Word
        logger.info('Finished.', 100)

        raise StopException
