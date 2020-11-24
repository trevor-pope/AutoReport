"""
Typographical related functions for manipulating and evaluating Word document elements.
"""
import logging
logger = logging.getLogger('email-automation')

from docx import Document
from docx.shared import Pt, RGBColor



def duplicate_run(paragraph, run):
    """
    Duplicate a run in a given paragraph. This creates a new run at the end of the paragraph that has the same
    text, font, and style.

    The run does not necessarily have to belong to the given paragraph.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        The paragraph that the new run is added onto
    run : docx.text.run.Run
        The run that will be duplicated

    Returns
    -------
    new_run : docx.text.run.Run
        The new duplicated run
    """
    unsettables = {'element', 'part', 'web_hidden', 'underline', 'superscript', 'subscript', 'strike',
                   'spec_vanish', 'snap_to_grid', 'small_caps'}
    new_run = paragraph.add_run(run.text)
    new_run.style = run.style
    for attr in dir(run.font):
        if not attr.startswith('_') and attr not in unsettables:
            if attr == 'color':  # For some reason, font.color is not settable, but font.color.rgb is
                new_run.font.color.rgb = run.font.color.rgb
                new_run.font.color.theme_color = run.font.color.theme_color
            else:
                setattr(new_run.font, attr, getattr(run.font, attr))

    return new_run


def is_default_typography(run):
    """
    Determines if the given run has the same typography as a default Document text.

    Parameters
    ----------
    run : docx.text.run.Run

    Returns
    -------
    bool
        True if the run is of default style, False otherwise
    """

    default = Document().add_paragraph().add_run('some default text')
    return are_same_typography(run, default)


def are_same_typography(run1, run2):
    """
    Determines whether two runs have the same typography (i.e. font, color, format, etc).

    Parameters
    ----------
    run1 : docx.text.run.Run
        The first run to compare
    run2 : docx.text.run.Run
        The second run to compare

    Returns
    -------
    bool :
        Whether or not the runs share the same typography
    """

    # Font attributes are None when inherited from the run's Style. Therefore, we first check for differences in Style,
    # and then check for differences in Font.
    style_attrs = ['style_id', 'name', 'type', 'priority', 'hidden', 'locked']
    style_attrs = ['style.' + attr for attr in style_attrs]

    font_attrs = ['name', 'bold', 'color.rgb', 'italic', 'size', 'cs_bold', 'cs_italic', 'double_strike', 'emboss',
                  'underline', 'hidden', 'highlight_color', 'imprint', 'no_proof', 'outline', 'rtl', 'all_caps',
                  'shadow', 'small_caps', 'snap_to_grid', 'spec_vanish', 'strike', 'subscript', 'superscript',
                  'web_hidden']
    font_attrs = ['font.' + attr for attr in font_attrs]

    for attr in style_attrs + font_attrs:
        if nested_getattr(run1, attr) != nested_getattr(run2, attr):
            return False

    return True


def nested_getattr(obj, attr):
    """
    Gets the value of an attribute, even if nested within other attributes via the dot operator.

    Parameters
    ----------
    obj
       The object whose attribute we are getting
    attr : str
        The name of the attribute and any parent attributes it descends from.
        For example, "attr1.attr2.attr3" would retrieve obj.attr1.attr2.attr3

    Returns
    -------
    value
        The value of the final attribute

    """
    attrs = attr.split('.')

    if len(attrs) > 1:
        return nested_getattr(getattr(obj, attrs[0]), '.'.join(attrs[1:]))
    else:
        return getattr(obj, attr)


def unite_similar_runs(runs):
    """
    Combine the text of runs that share the same typography and empty the left over runs.
    This modifies the paragraph in place, but only returns the non-empty runs out-of-place.

    Parameters
    ----------
    runs : List
        The runs that are being combined

    Returns
    -------
    non_empty_runs : List
        The list of non-emptied runs

    """
    if len(runs) <= 1:
        return runs

    last_run = None
    non_empty_runs = []

    for run in runs:
        if last_run is None:
            non_empty_runs.append(run)
            last_run = run
            continue

        if are_same_typography(run, last_run):
            last_run.text += run.text
            run.clear()
        else:
            non_empty_runs.append(run)
            last_run = run

    return non_empty_runs


def insert_runs(paragraph, runs, where, cuts=None):
    """
    Insert a list of runs inside a paragraph.

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        The paragraph runs are being inserted into.
    runs : List of docx.text.paragraph.Paragraph
        A list of runs to insert.
    where : int
        The index of the run within the given paragraph that will have other runs inserted.
    cuts : tuple of int, optional
        The indices of the text within the "where" run  that will be replaced by the new runs.
        Text between these indices will be erased. If None, runs will be inserted at the end of the specified run index.
    """
    initial_paragraph_length = len(paragraph.runs)
    if cuts:
        cut1, cut2 = cuts
        pre_cut_text, post_cut_text = paragraph.runs[where].text[:cut1], paragraph.runs[where].text[cut2:]
    else:
        pre_cut_text, post_cut_text = paragraph.runs[where].text, ''  # TODO test if passing in no cuts works

    # Create new empty runs to shift down old runs
    for i in range(len(runs) + 1):
        paragraph.add_run()

    # Shift down old runs
    j = -1
    for run in paragraph.runs[initial_paragraph_length-1:where:-1]:
        paragraph.runs[j].text = run.text
        set_typography_from_reference(paragraph.runs[j], run)
        j -= 1

    paragraph.runs[j].text = post_cut_text
    set_typography_from_reference(paragraph.runs[j], paragraph.runs[where])
    paragraph.runs[where].text = pre_cut_text

    # Insert new runs
    for i, run in enumerate(paragraph.runs[where + 1: where + len(runs) + 1]):
        run.text = runs[i].text
        set_typography_from_reference(run, runs[i])


def set_typography_from_reference(run, reference):
    """
    Sets a run's typography (i.e. font, color, format, etc) from a reference run.
    This changes the run in place.

    Parameters
    ----------
    run : docx.text.run.Run
        The run whose typography is being set
    reference : docx.text.run.Run
        The run whose typography is being copied
    """
    unsettables = {'element', 'part', 'web_hidden', 'underline', 'superscript', 'subscript', 'strike',
                   'spec_vanish', 'snap_to_grid', 'small_caps'}

    run.style = reference.style
    for attr in dir(reference.font):
        if not attr.startswith('_') and attr not in unsettables:
            if attr == 'color':
                run.font.color.rgb = reference.font.color.rgb
                # run.font.color.theme_color = reference.font.color.theme_color
            else:
                setattr(run.font, attr, getattr(reference.font, attr))


def create_floating_run(text, font_name, font_size, bold, italic, underline, color, **kwargs):
    """
    Create a new run in a floating paragraph.  TODO: Add kwargs to be added
    Parameters
    ----------
    text : str
        The run's text
    font_name : str
        The name of the font to be applied.
    font_size : int or float
        Pt size for the font.
    bold : bool
        Whether or not to apply bold to the run.
    italic : bool
        Whether or not to apply italics to the run.
    underline : bool
        Whether or not to apply underline to the run.
    color : tuple
        A 3-tuple containing the RGB of the color to apply to the run.

    Returns
    -------
    run : docx.text.run.Run
        The newly created run.
    """
    run = Document().add_paragraph().add_run(text)
    run.bold, run.italic, run.underline = bold, italic, underline
    run.font.name, run.font.size, run.font.color.rgb = font_name, Pt(font_size), RGBColor(*color)

    return run

